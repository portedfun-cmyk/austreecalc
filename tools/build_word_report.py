#!/usr/bin/env python3
"""Build a Word (.docx) report with embedded graphs from an AusTreeCalc JSON export.

Usage:
    python build_word_report.py path/to/export.json [output.docx]

The JSON file should be created from the Flutter app using the
"Export JSON for Word" button in the Report card.
"""

import json
import sys
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches


def _plot_sf_curve(x, y, xlabel, ylabel, title, out_path: Path):
    """Create a simple SF curve plot and save as PNG."""
    if not x or not y or len(x) != len(y):
        return None

    fig, ax = plt.subplots(figsize=(5, 3))
    ax.plot(x, y, marker="o")
    ax.set_title(title)
    ax.set_xlabel(xlabel)
    ax.set_ylabel(ylabel)
    ax.grid(True, linestyle="--", alpha=0.4)

    # Horizontal guides at SF = 1 and 1.5 to mirror the app visuals
    for level in (1.0, 1.5):
        ax.axhline(level, color="grey", linestyle=":", linewidth=0.8)

    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)
    return out_path


def build_word_report(json_path: Path, output_path: Path | None = None) -> Path:
    data = json.loads(json_path.read_text(encoding="utf-8"))

    tree = data.get("tree", {})
    inputs = data.get("inputs", {})
    defects = data.get("defects", {})
    decay = data.get("decay", {})
    graphs = data.get("graphs", {})
    text = data.get("text", {})

    doc = Document()

    # Title
    title = tree.get("label") or "AusTreeCalc tree stability report"
    doc.add_heading(title, level=0)

    # Tree and site info
    doc.add_heading("Tree and site details", level=1)
    p = doc.add_paragraph()
    p.add_run("Species: ").bold = True
    p.add_run(str(tree.get("species", "")))
    location = tree.get("site_location")
    if location:
        p = doc.add_paragraph()
        p.add_run("Location: ").bold = True
        p.add_run(str(location))

    doc.add_heading("Key inputs", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid"

    def _row(label: str, value) -> None:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = "" if value is None else str(value)

    _row("DBH (cm)", inputs.get("dbh_cm"))
    _row("Height (m)", inputs.get("height_m"))
    _row("Crown diameter (m)", inputs.get("crown_diameter_m"))
    _row("Cavity inner diameter (cm)", inputs.get("cavity_inner_diameter_cm"))
    _row("Design wind speed (m/s)", inputs.get("design_wind_speed_ms"))
    _row("Site factor", inputs.get("site_factor"))

    # Defects
    doc.add_heading("Observed structural defects / decay indicators", level=1)
    defect_list = doc.add_paragraph(style="List Bullet")
    flags = [
        ("Bracket fungi on stem or base", defects.get("bracket_fungi")),
        ("Cavity with visible decay", defects.get("cavity_decay")),
        ("Longitudinal cracks / shear planes", defects.get("cracks")),
        ("Basal/root-plate decay symptoms", defects.get("basal_decay")),
        ("Included bark / compromised unions", defects.get("union")),
    ]
    any_flag = False
    for label, flag in flags:
        if flag:
            defect_list = doc.add_paragraph(style="List Bullet")
            defect_list.add_run(label)
            any_flag = True
    other = defects.get("other")
    if other:
        defect_list = doc.add_paragraph(style="List Bullet")
        defect_list.add_run(str(other))
        any_flag = True
    if not any_flag:
        doc.add_paragraph("No specific structural defects selected.")

    k_defect = defects.get("strength_factor_k_defect")
    if k_defect is not None:
        para = doc.add_paragraph()
        para.add_run("Defect strength factor k_defect: ").bold = True
        para.add_run(f"{k_defect:.2f}")

    # Short summary and technical appendix text (already AS 4970-aligned)
    short_summary = text.get("short_summary", "").strip()
    if short_summary:
        doc.add_heading("Short summary (AS 4970-aligned)", level=1)
        for line in short_summary.splitlines():
            doc.add_paragraph(line)

    technical = text.get("technical_appendix", "").strip()
    if technical:
        doc.add_heading("Technical appendix", level=1)
        for line in technical.splitlines():
            doc.add_paragraph(line)

    base_calc = text.get("base_calculation", "").strip()
    if base_calc:
        doc.add_heading("Base calculation breakdown", level=1)
        for line in base_calc.splitlines():
            doc.add_paragraph(line)

    # Decay / residual wall info
    if decay:
        doc.add_heading("Decay / residual wall", level=1)
        cur = decay.get("current_residual_percent")
        crit = decay.get("critical_residual_percent")
        crit_wall = decay.get("critical_wall_thickness_cm")
        if cur is not None:
            doc.add_paragraph(
                f"Current residual wall (from DBH and cavity): {cur:.0f}% of diameter."
            )
        if crit is not None and crit_wall is not None:
            doc.add_paragraph(
                "At the current design wind speed, SF ≈ 1 when residual wall is "
                f"about {crit:.0f}% of diameter (≈ {crit_wall:.1f} cm on each side)."
            )

    # Graphs
    image_dir = output_path.parent if output_path is not None else json_path.parent
    image_dir.mkdir(parents=True, exist_ok=True)

    doc.add_heading("Graphs", level=1)

    fig_counter = 1

    # SF vs wind
    sf_wind = graphs.get("sf_vs_wind") or {}
    x = sf_wind.get("wind_ms") or []
    y = sf_wind.get("sf") or []
    if x and y:
        img_path = image_dir / "figure_sf_vs_wind.png"
        _plot_sf_curve(x, y, "Wind speed (m/s)", "SF", "SF vs wind speed", img_path)
        doc.add_paragraph(f"Figure {fig_counter} – Safety factor versus wind speed")
        doc.add_picture(str(img_path), width=Inches(5.0))
        fig_counter += 1

    # SF vs residual wall
    sf_rw = graphs.get("sf_vs_residual_wall") or {}
    x = sf_rw.get("residual_wall_percent") or []
    y = sf_rw.get("sf") or []
    if x and y:
        img_path = image_dir / "figure_sf_vs_residual_wall.png"
        _plot_sf_curve(
            x,
            y,
            "Residual wall (% of diameter)",
            "SF",
            "SF vs residual wall",
            img_path,
        )
        doc.add_paragraph(
            f"Figure {fig_counter} – Safety factor versus residual wall thickness"
        )
        doc.add_picture(str(img_path), width=Inches(5.0))
        fig_counter += 1

    # SF vs crown reduction
    sf_red = graphs.get("sf_vs_crown_reduction") or {}
    x = sf_red.get("reduction_percent") or []
    y = sf_red.get("sf") or []
    if x and y:
        img_path = image_dir / "figure_sf_vs_crown_reduction.png"
        _plot_sf_curve(
            x,
            y,
            "Crown reduction (%)",
            "SF",
            "SF vs crown reduction",
            img_path,
        )
        doc.add_paragraph(
            f"Figure {fig_counter} – Safety factor versus crown reduction (%)"
        )
        doc.add_picture(str(img_path), width=Inches(5.0))
        fig_counter += 1

    if output_path is None:
        output_path = json_path.with_suffix(".docx")

    doc.save(str(output_path))
    return output_path


def main(argv: list[str]) -> int:
    if len(argv) < 2:
        print(
            "Usage: python build_word_report.py path/to/export.json [output.docx]",
            file=sys.stderr,
        )
        return 1

    json_path = Path(argv[1]).expanduser().resolve()
    if not json_path.is_file():
        print(f"JSON file not found: {json_path}", file=sys.stderr)
        return 1

    output_path: Path | None = None
    if len(argv) >= 3:
        output_path = Path(argv[2]).expanduser().resolve()

    out = build_word_report(json_path, output_path)
    print(f"Wrote Word report to: {out}")
    return 0


if __name__ == "__main__":  # pragma: no cover
    raise SystemExit(main(sys.argv))
