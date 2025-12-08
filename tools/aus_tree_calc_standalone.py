#!/usr/bin/env python3
"""AusTreeCalc standalone Python program.

This script runs a simplified version of the AusTreeCalc statics engine
entirely in Python. It:

- Asks you for tree inputs via the terminal.
- Computes wind loading, bending stress and safety factor (SF).
- Models defects/decay via a strength reduction factor k_defect.
- Builds SF vs wind, SF vs residual wall and SF vs crown-reduction curves.
- Generates a Word (.docx) report with tables, text and embedded graphs.

Usage (from project root):

    pip3 install -r requirements.txt   # one-time
    python3 tools/aus_tree_calc_standalone.py

The report will be written next to this script as
`aus_tree_calc_report.docx` by default.
"""

from __future__ import annotations

import json
import math
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import List, Tuple

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

AIR_DENSITY = 1.2  # kg/m3


@dataclass
class SpeciesPreset:
    id: str
    name: str
    fb_green_mpa: float
    drag_coefficient: float
    crown_shape_factor: float
    default_fullness: float


SPECIES_PRESETS: List[SpeciesPreset] = [
    SpeciesPreset(
        id="euc_high",
        name="Eucalypt – High Strength (ironbark / spotted gum type)",
        fb_green_mpa=50.0,
        drag_coefficient=0.25,
        crown_shape_factor=0.7,
        default_fullness=0.9,
    ),
    SpeciesPreset(
        id="euc_typical",
        name="Eucalypt – Typical Street Tree",
        fb_green_mpa=35.0,
        drag_coefficient=0.25,
        crown_shape_factor=0.7,
        default_fullness=0.9,
    ),
    SpeciesPreset(
        id="broadleaf_deciduous",
        name="Broadleaf – Plane / Elm / Oak",
        fb_green_mpa=28.0,
        drag_coefficient=0.30,
        crown_shape_factor=0.75,
        default_fullness=0.95,
    ),
    SpeciesPreset(
        id="conifer_softwood",
        name="Conifer – Pine / Cypress",
        fb_green_mpa=20.0,
        drag_coefficient=0.35,
        crown_shape_factor=0.8,
        default_fullness=1.0,
    ),
    SpeciesPreset(
        id="araucaria",
        name="Araucaria – Norfolk Island Pine",
        fb_green_mpa=24.0,
        drag_coefficient=0.30,
        crown_shape_factor=0.7,
        default_fullness=0.95,
    ),
    SpeciesPreset(
        id="unknown_hardwood",
        name="Unknown Hardwood (broadleaf)",
        fb_green_mpa=25.0,
        drag_coefficient=0.28,
        crown_shape_factor=0.7,
        default_fullness=0.9,
    ),
    SpeciesPreset(
        id="unknown_softwood",
        name="Unknown Softwood / Evergreen",
        fb_green_mpa=18.0,
        drag_coefficient=0.33,
        crown_shape_factor=0.75,
        default_fullness=0.95,
    ),
]


@dataclass
class CalcResult:
    safety_factor: float
    bending_stress_mpa: float
    q_pa: float
    wind_force_n: float
    bending_moment_nm: float


def prompt_float(prompt: str, default: float | None = None) -> float:
    while True:
        txt = input(f"{prompt} " + (f"[{default}] " if default is not None else ""))
        txt = txt.strip()
        if not txt and default is not None:
            return float(default)
        try:
            value = float(txt.replace(",", "."))
        except ValueError:
            print("Please enter a number.")
            continue
        if value <= 0:
            print("Please enter a value greater than zero.")
            continue
        return value


def prompt_optional_float(prompt: str) -> float | None:
    txt = input(f"{prompt} [blank for none] ").strip()
    if not txt:
        return None
    try:
        value = float(txt.replace(",", "."))
    except ValueError:
        print("Could not parse number, treating as none.")
        return None
    if value <= 0:
        return None
    return value


def prompt_yes_no(prompt: str, default: bool = False) -> bool:
    suffix = "[Y/n] " if default else "[y/N] "
    while True:
        txt = input(f"{prompt} {suffix}").strip().lower()
        if not txt:
            return default
        if txt in {"y", "yes"}:
            return True
        if txt in {"n", "no"}:
            return False
        print("Please answer y or n.")


def choose_species() -> SpeciesPreset:
    print("Select species / strength group:")
    for i, sp in enumerate(SPECIES_PRESETS, start=1):
        print(f"  {i}. {sp.name}")
    while True:
        txt = input(f"Enter number [2 for '{SPECIES_PRESETS[1].name}']: ").strip()
        if not txt:
            return SPECIES_PRESETS[1]
        try:
            idx = int(txt)
        except ValueError:
            print("Please enter a number from the list.")
            continue
        if 1 <= idx <= len(SPECIES_PRESETS):
            return SPECIES_PRESETS[idx - 1]
        print("Please choose a valid number.")


def compute_defect_strength_factor(
    bracket_fungi: bool,
    cavity_decay: bool,
    cracks: bool,
    basal_decay: bool,
    union: bool,
) -> float:
    k = 1.0
    if bracket_fungi:
        k *= 0.8
    if cavity_decay:
        k *= 0.8
    if cracks:
        k *= 0.9
    if basal_decay:
        k *= 0.8
    if union:
        k *= 0.9
    if k < 0.3:
        k = 0.3
    if k > 1.0:
        k = 1.0
    return k


def calculate_single(
    species: SpeciesPreset,
    dbh_cm: float,
    height_m: float,
    crown_diameter_m: float,
    design_wind_ms: float,
    cavity_inner_cm: float | None,
    fullness_override: float | None,
    site_factor: float,
    k_defect: float,
) -> CalcResult:
    dbh_m = dbh_cm / 100.0
    d_outer = dbh_m
    d_inner = 0.0
    if cavity_inner_cm is not None and cavity_inner_cm > 0:
        cav = cavity_inner_cm
        if cav >= dbh_cm:
            cav = dbh_cm * 0.99
        d_inner = cav / 100.0

    V = design_wind_ms
    q = site_factor * 0.5 * AIR_DENSITY * V * V

    radius_crown = crown_diameter_m / 2.0
    a_plan = math.pi * radius_crown * radius_crown

    fullness_base = fullness_override if fullness_override is not None else species.default_fullness
    fullness = max(0.1, min(1.0, fullness_base))
    area = a_plan * species.crown_shape_factor * fullness

    wind_force = q * species.drag_coefficient * area
    h_eff = 0.66 * height_m
    m_wind = wind_force * h_eff

    if d_inner > 0.0:
        W = math.pi * (d_outer**4 - d_inner**4) / (32.0 * d_outer)
    else:
        W = math.pi * d_outer**3 / 32.0

    sigma_pa = m_wind / W
    sigma_mpa = sigma_pa / 1e6

    effective_fb = species.fb_green_mpa * k_defect
    sf = effective_fb / sigma_mpa if sigma_mpa > 0 else float("inf")

    return CalcResult(
        safety_factor=sf,
        bending_stress_mpa=sigma_mpa,
        q_pa=q,
        wind_force_n=wind_force,
        bending_moment_nm=m_wind,
    )


def estimate_wind_to_failure(
    species: SpeciesPreset,
    dbh_cm: float,
    height_m: float,
    crown_diameter_m: float,
    design_wind_ms: float,
    cavity_inner_cm: float | None,
    fullness_override: float | None,
    site_factor: float,
    k_defect: float,
) -> float | None:
    ref = calculate_single(
        species,
        dbh_cm,
        height_m,
        crown_diameter_m,
        design_wind_ms,
        cavity_inner_cm,
        fullness_override,
        site_factor,
        k_defect,
    )
    sf = ref.safety_factor
    if not math.isfinite(sf) or sf <= 0:
        return None
    return design_wind_ms * math.sqrt(sf)


def residual_wall_fraction(dbh_cm: float, cavity_inner_cm: float | None) -> float:
    if dbh_cm <= 0:
        return 1.0
    if cavity_inner_cm is None or cavity_inner_cm <= 0:
        return 1.0
    cav = cavity_inner_cm
    if cav >= dbh_cm:
        cav = dbh_cm * 0.99
    frac = (dbh_cm - cav) / dbh_cm
    return max(0.0, min(1.0, frac))


def build_sf_vs_wind_curve(
    species: SpeciesPreset,
    dbh_cm: float,
    height_m: float,
    crown_diameter_m: float,
    design_wind_ms: float,
    cavity_inner_cm: float | None,
    fullness_override: float | None,
    site_factor: float,
    k_defect: float,
    wind_to_failure: float | None,
) -> Tuple[List[float], List[float]]:
    if design_wind_ms <= 0:
        return [], []
    min_v = max(5.0, design_wind_ms * 0.5)
    max_v = design_wind_ms * 1.8
    if wind_to_failure is not None and math.isfinite(wind_to_failure) and wind_to_failure > 0:
        extended = wind_to_failure * 1.1
        if extended > max_v:
            max_v = extended
    if max_v <= min_v:
        max_v = min_v + 5.0
    xs: List[float] = []
    ys: List[float] = []
    steps = 12
    for i in range(steps):
        v = min_v + (max_v - min_v) * i / (steps - 1)
        res = calculate_single(
            species,
            dbh_cm,
            height_m,
            crown_diameter_m,
            v,
            cavity_inner_cm,
            fullness_override,
            site_factor,
            k_defect,
        )
        xs.append(v)
        ys.append(res.safety_factor)
    return xs, ys


def build_sf_vs_residual_wall_curve(
    species: SpeciesPreset,
    dbh_cm: float,
    height_m: float,
    crown_diameter_m: float,
    design_wind_ms: float,
    fullness_override: float | None,
    site_factor: float,
    k_defect: float,
) -> Tuple[List[float], List[float], float | None, float | None]:
    xs: List[float] = []
    ys: List[float] = []
    rw_min, rw_max = 20.0, 100.0
    steps = 9
    for i in range(steps):
        rw = rw_min + (rw_max - rw_min) * i / (steps - 1)
        frac = rw / 100.0
        cav_cm = dbh_cm * (1.0 - frac)
        cavity_sim = cav_cm if cav_cm > 0 else None
        res = calculate_single(
            species,
            dbh_cm,
            height_m,
            crown_diameter_m,
            design_wind_ms,
            cavity_sim,
            fullness_override,
            site_factor,
            k_defect,
        )
        xs.append(rw)
        ys.append(res.safety_factor)

    critical_rw: float | None = None
    if len(xs) >= 2:
        for i in range(len(xs) - 1):
            y1, y2 = ys[i], ys[i + 1]
            if not (math.isfinite(y1) and math.isfinite(y2)):
                continue
            if (y1 >= 1.0 and y2 <= 1.0) or (y1 <= 1.0 and y2 >= 1.0):
                x1, x2 = xs[i], xs[i + 1]
                t = (1.0 - y1) / (y2 - y1) if y2 != y1 else 0.0
                x = x1 + (x2 - x1) * t
                critical_rw = max(rw_min, min(rw_max, x))
                break

    critical_wall_cm: float | None = None
    if critical_rw is not None:
        critical_wall_cm = dbh_cm * (critical_rw / 100.0) / 2.0

    return xs, ys, critical_rw, critical_wall_cm


def build_sf_vs_crown_reduction_curve(
    species: SpeciesPreset,
    dbh_cm: float,
    height_m: float,
    crown_diameter_m: float,
    design_wind_ms: float,
    cavity_inner_cm: float | None,
    fullness_override: float | None,
    site_factor: float,
    k_defect: float,
    base_reduction_percent: float,
    fullness_reduction_percent: float,
) -> Tuple[List[float], List[float]]:
    xs: List[float] = []
    ys: List[float] = []
    max_red = max(5.0, min(40.0, base_reduction_percent or 20.0))
    steps = 9
    fullness_base = (
        fullness_override if fullness_override is not None else species.default_fullness
    )
    for i in range(steps):
        r = max_red * i / (steps - 1)
        crown_after = crown_diameter_m * (1.0 - r / 100.0)
        fullness_after_raw = fullness_base * (1.0 - fullness_reduction_percent / 100.0)
        fullness_after = max(0.1, min(1.0, fullness_after_raw))
        res = calculate_single(
            species,
            dbh_cm,
            height_m,
            crown_after,
            design_wind_ms,
            cavity_inner_cm,
            fullness_after,
            site_factor,
            k_defect,
        )
        xs.append(r)
        ys.append(res.safety_factor)
    return xs, ys


def plot_sf_curve(x, y, xlabel: str, ylabel: str, title: str, out_path: Path) -> None:
    if not x or not y or len(x) != len(y):
        return
    fig, ax = plt.subplots(figsize=(5, 3))
    ax.plot(x, y, marker="o")
    ax.set_title(title)
    ax.set_xlabel(xlabel)
    ax.set_ylabel(ylabel)
    ax.grid(True, linestyle="--", alpha=0.4)
    for level in (1.0, 1.5):
        ax.axhline(level, color="grey", linestyle=":", linewidth=0.8)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)


def build_word_report_from_python(
    output_path: Path,
    tree_label: str,
    species: SpeciesPreset,
    site_location: str,
    inputs: dict,
    defects: dict,
    result: CalcResult,
    wind_to_failure: float | None,
    decay_info: dict,
    graphs: dict,
) -> None:
    doc = Document()

    doc.add_heading(tree_label or "AusTreeCalc tree stability report", level=0)

    # Tree and site
    doc.add_heading("Tree and site details", level=1)
    p = doc.add_paragraph()
    p.add_run("Species: ").bold = True
    p.add_run(species.name)
    if site_location:
        p = doc.add_paragraph()
        p.add_run("Location: ").bold = True
        p.add_run(site_location)

    # Inputs table
    doc.add_heading("Key inputs", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid"

    def row(label: str, value) -> None:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = "" if value is None else str(value)

    row("DBH (cm)", inputs.get("dbh_cm"))
    row("Height (m)", inputs.get("height_m"))
    row("Crown diameter (m)", inputs.get("crown_diameter_m"))
    row("Cavity inner diameter (cm)", inputs.get("cavity_inner_diameter_cm"))
    row("Design wind speed (m/s)", inputs.get("design_wind_speed_ms"))
    row("Site factor", inputs.get("site_factor"))

    # Defects
    doc.add_heading("Observed structural defects / decay indicators", level=1)
    flags = [
        ("Bracket fungi on stem or base", defects.get("bracket_fungi")),
        ("Cavity with visible decay", defects.get("cavity_decay")),
        ("Longitudinal cracks / shear planes", defects.get("cracks")),
        ("Basal/root-plate decay symptoms", defects.get("basal_decay")),
        ("Included bark / compromised unions", defects.get("union")),
    ]
    any_flag = False
    for label, flag in flags:
        if flag:
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(label)
            any_flag = True
    other = defects.get("other")
    if other:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(str(other))
        any_flag = True
    if not any_flag:
        doc.add_paragraph("No specific structural defects selected.")

    k_defect = defects.get("strength_factor_k_defect")
    if k_defect is not None:
        para = doc.add_paragraph()
        para.add_run("Defect strength factor k_defect: ").bold = True
        para.add_run(f"{k_defect:.2f}")

    # Numerical results
    doc.add_heading("Numerical results", level=1)
    para = doc.add_paragraph()
    para.add_run("Safety factor at design wind speed (SF): ").bold = True
    para.add_run(f"{result.safety_factor:.2f}" if math.isfinite(result.safety_factor) else "∞")

    para = doc.add_paragraph()
    para.add_run("Bending stress at governing section: ").bold = True
    para.add_run(f"{result.bending_stress_mpa:.2f} MPa")

    if wind_to_failure is not None and math.isfinite(wind_to_failure):
        para = doc.add_paragraph()
        para.add_run("Estimated wind-to-failure speed (SF ≈ 1): ").bold = True
        para.add_run(f"{wind_to_failure:.1f} m/s")

    # Decay / residual wall
    doc.add_heading("Decay / residual wall", level=1)
    cur = decay_info.get("current_residual_percent")
    crit = decay_info.get("critical_residual_percent")
    crit_wall = decay_info.get("critical_wall_thickness_cm")
    if cur is not None:
        doc.add_paragraph(
            f"Current residual wall (from DBH and cavity): {cur:.0f}% of diameter.")
    if crit is not None and crit_wall is not None:
        doc.add_paragraph(
            "At the current design wind speed, SF ≈ 1 when residual wall is "
            f"about {crit:.0f}% of diameter (≈ {crit_wall:.1f} cm on each side)."
        )

    # Graphs
    doc.add_heading("Graphs", level=1)
    image_dir = output_path.parent
    image_dir.mkdir(parents=True, exist_ok=True)

    fig_counter = 1

    def add_curve_graph(key: str, xlabel: str, title: str, filename: str) -> None:
        nonlocal fig_counter
        g = graphs.get(key) or {}
        xs = g.get("x") or []
        ys = g.get("y") or []
        if not xs or not ys:
            return
        img_path = image_dir / filename
        plot_sf_curve(xs, ys, xlabel, "SF", title, img_path)
        doc.add_paragraph(
            f"Figure {fig_counter} – {title}")
        doc.add_picture(str(img_path), width=Inches(5.0))
        fig_counter += 1

    add_curve_graph("sf_vs_wind", "Wind speed (m/s)", "Safety factor versus wind speed",
                    "figure_sf_vs_wind.png")
    add_curve_graph(
        "sf_vs_residual_wall",
        "Residual wall (% of diameter)",
        "Safety factor versus residual wall thickness",
        "figure_sf_vs_residual_wall.png",
    )
    add_curve_graph(
        "sf_vs_crown_reduction",
        "Crown reduction (%)",
        "Safety factor versus crown reduction (%)",
        "figure_sf_vs_crown_reduction.png",
    )

    doc.save(str(output_path))


def main(argv: List[str]) -> int:
    print("AusTreeCalc standalone (Python)")
    print("This will ask you for tree details and generate a Word report.\n")

    tree_label = input("Tree label / ID [Tree 1]: ").strip() or "Tree 1"
    site_location = input("Site / location [optional]: ").strip()

    species = choose_species()

    dbh_cm = prompt_float("DBH (cm):", 50.0)
    height_m = prompt_float("Tree height (m):", 18.0)
    crown_diameter_m = prompt_float("Crown diameter (m):", 10.0)
    cavity_inner_cm = prompt_optional_float("Cavity inner diameter (cm)")

    design_wind_ms = prompt_float("Design wind speed (m/s):", 40.0)
    site_factor = prompt_float("Site factor (exposure/topography) 0.5–1.5:", 1.0)

    use_fullness_override = prompt_yes_no(
        "Override crown fullness (default is species typical)?", False
    )
    fullness_override = None
    if use_fullness_override:
        fullness_override = prompt_float("Crown fullness (0–1):", species.default_fullness)
        if fullness_override < 0.1:
            fullness_override = 0.1
        if fullness_override > 1.0:
            fullness_override = 1.0

    print("\nObserved structural defects / decay indicators:")
    defect_bracket = prompt_yes_no("  Bracket fungi on stem or base?", False)
    defect_cavity_decay = prompt_yes_no("  Cavity with visible decay?", False)
    defect_cracks = prompt_yes_no("  Longitudinal cracks / shear planes?", False)
    defect_basal = prompt_yes_no("  Basal/root-plate decay symptoms?", False)
    defect_union = prompt_yes_no("  Included bark / compromised unions?", False)
    defect_other = input("  Other defect / decay observations (optional): ").strip()

    k_defect = compute_defect_strength_factor(
        defect_bracket,
        defect_cavity_decay,
        defect_cracks,
        defect_basal,
        defect_union,
    )

    # Main calculation
    result = calculate_single(
        species,
        dbh_cm,
        height_m,
        crown_diameter_m,
        design_wind_ms,
        cavity_inner_cm,
        fullness_override,
        site_factor,
        k_defect,
    )

    wind_to_failure = estimate_wind_to_failure(
        species,
        dbh_cm,
        height_m,
        crown_diameter_m,
        design_wind_ms,
        cavity_inner_cm,
        fullness_override,
        site_factor,
        k_defect,
    )

    # Curves
    sf_wind_x, sf_wind_y = build_sf_vs_wind_curve(
        species,
        dbh_cm,
        height_m,
        crown_diameter_m,
        design_wind_ms,
        cavity_inner_cm,
        fullness_override,
        site_factor,
        k_defect,
        wind_to_failure,
    )

    rw_x, rw_y, crit_rw, crit_wall = build_sf_vs_residual_wall_curve(
        species,
        dbh_cm,
        height_m,
        crown_diameter_m,
        design_wind_ms,
        fullness_override,
        site_factor,
        k_defect,
    )

    crown_red_base = prompt_float("Typical crown reduction to model (%)", 20.0)
    fullness_red = prompt_float(
        "Typical crown thinning effect on fullness (%)", 30.0
    )

    red_x, red_y = build_sf_vs_crown_reduction_curve(
        species,
        dbh_cm,
        height_m,
        crown_diameter_m,
        design_wind_ms,
        cavity_inner_cm,
        fullness_override,
        site_factor,
        k_defect,
        crown_red_base,
        fullness_red,
    )

    # Simple on-screen summary
    print("\nResults:")
    print(f"  Safety factor SF at design wind: {result.safety_factor:.2f}")
    print(f"  Bending stress: {result.bending_stress_mpa:.2f} MPa")
    if wind_to_failure is not None and math.isfinite(wind_to_failure):
        print(f"  Estimated wind-to-failure (SF ≈ 1): {wind_to_failure:.1f} m/s")

    res_wall_frac = residual_wall_fraction(dbh_cm, cavity_inner_cm)
    res_wall_pct = res_wall_frac * 100.0
    print(f"  Current residual wall (from DBH and cavity): {res_wall_pct:.0f}% of diameter")
    if crit_rw is not None and crit_wall is not None:
        print(
            "  SF ≈ 1 at residual wall ≈ "
            f"{crit_rw:.0f}% (≈ {crit_wall:.1f} cm on each side) at design wind."
        )

    # Build Word report with graphs
    script_dir = Path(__file__).resolve().parent
    output_doc = script_dir / "aus_tree_calc_report.docx"

    inputs = {
        "dbh_cm": dbh_cm,
        "height_m": height_m,
        "crown_diameter_m": crown_diameter_m,
        "cavity_inner_diameter_cm": cavity_inner_cm,
        "design_wind_speed_ms": design_wind_ms,
        "site_factor": site_factor,
    }
    defects = {
        "bracket_fungi": defect_bracket,
        "cavity_decay": defect_cavity_decay,
        "cracks": defect_cracks,
        "basal_decay": defect_basal,
        "union": defect_union,
        "other": defect_other,
        "strength_factor_k_defect": k_defect,
    }
    decay_info = {
        "current_residual_percent": res_wall_pct,
        "critical_residual_percent": crit_rw,
        "critical_wall_thickness_cm": crit_wall,
    }
    graphs = {
        "sf_vs_wind": {"x": sf_wind_x, "y": sf_wind_y},
        "sf_vs_residual_wall": {"x": rw_x, "y": rw_y},
        "sf_vs_crown_reduction": {"x": red_x, "y": red_y},
    }

    build_word_report_from_python(
        output_doc,
        tree_label,
        species,
        site_location,
        inputs,
        defects,
        result,
        wind_to_failure,
        decay_info,
        graphs,
    )

    print(f"\nWord report written to: {output_doc}")
    print("You can open this .docx in Word or Pages and edit as needed.")

    # Also save raw data as JSON next to the report for reference
    json_path = output_doc.with_suffix(".json")
    payload = {
        "tree_label": tree_label,
        "species_id": species.id,
        "species_name": species.name,
        "site_location": site_location,
        "inputs": inputs,
        "defects": defects,
        "result": {
            "safety_factor": result.safety_factor,
            "bending_stress_mpa": result.bending_stress_mpa,
            "q_pa": result.q_pa,
            "wind_force_n": result.wind_force_n,
            "bending_moment_nm": result.bending_moment_nm,
            "wind_to_failure_ms": wind_to_failure,
        },
        "decay": decay_info,
        "graphs": graphs,
    }
    json_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    print(f"Raw calculation data saved to: {json_path}")

    return 0


if __name__ == "__main__":  # pragma: no cover
    raise SystemExit(main(sys.argv))
